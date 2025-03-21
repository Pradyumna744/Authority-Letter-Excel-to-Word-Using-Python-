import os
import pandas as pd
import datetime
from docx import Document
from tkinter import Tk, Text, filedialog, Scrollbar
from tkinter.ttk import Progressbar, Label, Button, Entry, Frame
import ttkbootstrap as ttk
from ttkbootstrap.dialogs import Messagebox


def mail_merge(template_path, excel_path, output_folder, log_text, progress_bar):
    try:
        # Load the Excel file
        df = pd.read_excel(excel_path)
        df.columns = df.columns.str.strip()  # Remove extra spaces from column names

        total_records = len(df)
        log_text.insert("end", f"Loaded {total_records} records from the Excel file.\n")
        progress_bar["maximum"] = total_records

        # Get the current date
        current_date = datetime.datetime.now().strftime("%d-%m-%Y")

        # Add the current date to the DataFrame for placeholders
        df["Current Date"] = current_date

        # Identify columns with datetime data
        datetime_columns = df.select_dtypes(include=["datetime64", "datetime"]).columns
        for col in datetime_columns:
            df[col] = df[col].dt.strftime("%d-%m-%Y")

        # Process each row in the Excel file
        for index, row in df.iterrows():
            insurance_company = str(row.get("Insurance Company", "Unknown")).replace("/", "_").strip()
            policy_no = str(row.get("Policy No.", "")).replace("/", "_").strip()
            insured_name = str(row.get("Insured Name", "")).replace("/", "_").strip()

            # Create folders
            company_folder_path = os.path.join(output_folder, insurance_company)
            os.makedirs(company_folder_path, exist_ok=True)

            record_folder_name = f"{policy_no}_{insured_name}"
            record_folder_path = os.path.join(company_folder_path, record_folder_name)
            os.makedirs(record_folder_path, exist_ok=True)

            # Load and process Word template
            doc = Document(template_path)
            for paragraph in doc.paragraphs:
                for key in df.columns:
                    placeholder = f"[{key}]"
                    value = str(row.get(key, "")).strip()
                    if pd.isna(value) or value == "###" or value is None:
                        value = ""
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        for key in df.columns:
                            placeholder = f"[{key}]"
                            value = str(row.get(key, "")).strip()
                            if pd.isna(value) or value == "###" or value is None:
                                value = ""
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)

            # Replace the current date placeholder
            for paragraph in doc.paragraphs:
                if "[Current Date]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Current Date]", current_date)

            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        if "[Current Date]" in cell.text:
                            cell.text = cell.text.replace("[Current Date]", current_date)

            # Save the output document
            output_file = os.path.join(record_folder_path, f"Document_{policy_no}_{insured_name}.docx")
            doc.save(output_file)

            # Update log and progress bar
            log_text.insert("end", f"Processed: {policy_no} - {insured_name}\n")
            progress_bar["value"] = index + 1
            log_text.see("end")

        log_text.insert("end", f"Mail merge completed! Files saved in: {output_folder}\n")
        Messagebox.show_info("Success", "Mail merge completed successfully!", title="Operation Complete")

    except Exception as e:
        Messagebox.show_error("Error", f"An error occurred: {e}", title="Error")
        log_text.insert("end", f"Error: {e}\n")


def select_template():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if file_path:
        template_entry.delete(0, "end")
        template_entry.insert(0, file_path)


def select_excel():
    file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if file_path:
        excel_entry.delete(0, "end")
        excel_entry.insert(0, file_path)


def select_output_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        output_entry.delete(0, "end")
        output_entry.insert(0, folder_path)


def start_merge():
    template_path = template_entry.get()
    excel_path = excel_entry.get()
    output_folder = output_entry.get()

    if not template_path or not excel_path or not output_folder:
        Messagebox.show_warning("Missing Input", "Please provide all required inputs.", title="Warning")
        return

    log_text.delete(1.0, "end")  # Clear previous log
    mail_merge(template_path, excel_path, output_folder, log_text, progress_bar)


# Tkinter GUI Setup
root = ttk.Window(themename="journal")
root.title("Advanced Excel to Word Mail Merge")
root.geometry("800x600")
root.resizable(True, True)

# Create a Frame for input fields and buttons
input_frame = ttk.Frame(root, padding=20)
input_frame.grid(row=0, column=0, columnspan=3, sticky="ew")

# Input Fields
Label(input_frame, text="Word Template:").grid(row=0, column=0, sticky="w", padx=20, pady=10)
template_entry = Entry(input_frame, width=60)
template_entry.grid(row=0, column=1, padx=10)
Button(input_frame, text="Browse", command=select_template).grid(row=0, column=2, padx=10)

Label(input_frame, text="Excel File:").grid(row=1, column=0, sticky="w", padx=20, pady=10)
excel_entry = Entry(input_frame, width=60)
excel_entry.grid(row=1, column=1, padx=10)
Button(input_frame, text="Browse", command=select_excel).grid(row=1, column=2, padx=10)

Label(input_frame, text="Output Folder:").grid(row=2, column=0, sticky="w", padx=20, pady=10)
output_entry = Entry(input_frame, width=60)
output_entry.grid(row=2, column=1, padx=10)
Button(input_frame, text="Browse", command=select_output_folder).grid(row=2, column=2, padx=10)

# Progress Bar
progress_bar = Progressbar(root, orient="horizontal", mode="determinate", bootstyle="success-striped")
progress_bar.grid(row=3, column=0, columnspan=3, pady=10)

# Log Textbox
Label(root, text="Log:").grid(row=4, column=0, sticky="nw", padx=20, pady=5)
log_text = Text(root, height=12, width=80, wrap="word")
log_text.grid(row=5, column=0, columnspan=3, padx=20, pady=5)

scrollbar = Scrollbar(root, command=log_text.yview)
scrollbar.grid(row=5, column=3, sticky="ns")
log_text.config(yscrollcommand=scrollbar.set)

# Start Button
Button(root, text="Start Mail Merge", command=start_merge, bootstyle="primary").grid(row=6, column=1, pady=30)

# Start the Tkinter event loop
root.mainloop()
